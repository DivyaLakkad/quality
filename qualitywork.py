# To add a new cell, type '# %%'
# To add a new markdown cell, type '# %% [markdown]'
# %%
import pandas as pd
import numpy
numpy.datetime_as_string
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches
import os
from datetime import date, timedelta
from email.message import EmailMessage

import smtplib 
import base64
import email
import imaplib
from email.mime.multipart import MIMEMultipart 
from email.mime.text import MIMEText 
from email.mime.base import MIMEBase 
from email import encoders 
import email

def email_auto_script(to_email, subject, body, file_path):
    print('Sending Email!!!')
    print (file_path)

    gmail_user = "graham.scripting@gmail.com"
    gmail_password = "directorchris"

    
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = gmail_user
    msg["To"] = to_email
    msg.set_content(body)

    for file in file_path:
        print(file)
        with open(file, "rb") as f:
            file_data = f.read()
            file_name = os.path.split(file)
            file_name = file_name[-1]

            msg.add_attachment(
                file_data,
                maintype="application",
                subtype="octet-stream",
                filename=file_name,
            )

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login(gmail_user, gmail_password)
        smtp.send_message(msg)
        print("Email sent!")


new_list = []
new_df = pd.read_excel('NEW_FILE.xlsx', engine="openpyxl")
new_df.head(5)


# %%
old_df = pd.read_excel('OLD_FILE.xlsx',engine="openpyxl")


# %%
old_df.head(5)


# %%
diff_df = pd.concat([new_df,old_df]).drop_duplicates(keep=False)
diff_df.head()


# %%
new_df.reset_index(drop=True)


# %%
car = ['Project No.','Project Description','Report Date','Name','Occurrence No.','Description.']
fsr = ['Project No.','Project Description','Report Date','Name','Occurrence No.','Description.']
ncr = ['Project No.','Project Description','Occurrence No.','Description.']
par = ['Project No.','Project Description','Report Date','Name','Occurrence No.','Description.']
qsr = ['Project No.','Project Description','Report Date','Name','Occurrence No.','Type','Discipline','Description.']


# %%
CAR= [diff_df[label].values[0] for label in car if label in diff_df]
FSR= [diff_df[label].values[0] for label in fsr if label in diff_df]
NCR= [diff_df[label].values[0] for label in ncr if label in diff_df]
PAR= [diff_df[label].values[0] for label in par if label in diff_df]
QSR= [diff_df[label].values[0] for label in qsr if label in diff_df]
print(CAR)


# %%
import docx

def getText(filename):
    doc = docx.Document(filename)   
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


# %%
i = diff_df.shape[0]
for x in range(i):
    print(x)


# %%

from docx import Document
from docx2pdf import convert

files_email_path =[]
i = diff_df.shape[0]
for x in range(i):
    pre = '{}'
    CAR1= [diff_df[label].values[x] for label in car if label in diff_df]
    FSR1= [diff_df[label].values[x] for label in fsr if label in diff_df]
    NCR1= [diff_df[label].values[x] for label in ncr if label in diff_df]
    PAR1= [diff_df[label].values[x] for label in par if label in diff_df]
    QSR1= [diff_df[label].values[x] for label in qsr if label in diff_df]

    file_name= diff_df['Type of occurrence:'].values[x]
    document = Document(file_name+'.docx')

    if file_name == 'CAR':
        output = [diff_df[label].values[x] for label in car if label in diff_df]
    elif file_name == 'FSR':
        output = [diff_df[label].values[x] for label in fsr if label in diff_df]
    elif file_name == 'NCR':
        output = [diff_df[label].values[x] for label in ncr if label in diff_df]
    elif file_name == 'PAR':
        output = [diff_df[label].values[x] for label in par if label in diff_df]
    else:
        output = [diff_df[label].values[x] for label in qsr if label in diff_df]

    length = len(output)
    print(length)
    tables = document.tables

    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs: 
                    if len(output) != 0:
                        to = str(output[0])
                        if pre in paragraph.text:
                            paragraph.text = paragraph.text.replace(pre, to)
                            output.pop(0)
                            print(paragraph.text)
                        
                    else:
                        to = ''
                        paragraph.text = paragraph.text.replace(pre, to)


    newfilename =  str(diff_df['Project No.'].values[x]) + '.docx'
    
    
    document.save(newfilename) 
    convert(newfilename)
    newfilename =  str(diff_df['Project No.'].values[x]) + '.pdf'
    print(newfilename)
    newfilepath =  os.path.join(os.path.dirname(__file__), newfilename)
    newfilepath= [r"{}".format(newfilepath)]
    print("this is the one")
    print(newfilepath)
    #Change email_to to match the correct column from  NEW_FILE
    #strftime
    #email_to = [diff_df['Email'].values[x]]
    email_to = ['divya.lakkad@graham.ca','basharat.javed@graham.ca']
    subject = 'Test 1'
    body = "Hey Divya!\nThis is a test to see if your function works!!\n\nRegards\nGraham Scripting"
    

    email_auto_script(email_to, subject, body, newfilepath)

    #convert("test.docx","output.pdf") 


# %%




# %%
# file_path = []
# email_to = ['divya.lakkad@graham.ca','momin.ejaz@graham.ca','basharat.javed@graham.ca']
# subject = 'Test 1'
# body = "Hey Divya!\nThis is a test to see if your function works!!\n\nRegards\nGraham Scripting"
# file_path = [r'C:\Users\divyal\Desktop\projects\xl\abc.pdf']

# email_auto_script(email_to, subject, body, file_path)

