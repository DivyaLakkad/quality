# write try and except for teh code
import time
import sched
from email import encoders
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import smtplib
import autopep8
import shutil
import docx
from docx2pdf import convert
import xlsxwriter
import imaplib
import email
import base64
from email.message import EmailMessage
from datetime import datetime
import datetime
from datetime import date, timedelta
import os
from docx.shared import Inches
from docx import Document
from openpyxl import load_workbook
import pandas as pd
import numpy
numpy.datetime_as_string

s = sched.scheduler(time.time, time.sleep)


def email_auto_script(to_email, subject, body, file_path):
    print('Sending Email!!!')

    gmail_user = "graham.scripting@gmail.com"
    gmail_password = "directorchris"

    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = gmail_user
    msg["To"] = to_email
    #msg["Bcc"] = 'divya.lakkad@graham.ca'
    msg.set_content(body)

    for file in file_path:
        print(file)
        with open(file, "rb") as f:
            file_data = f.read()
            file_name = os.path.split(file)
            file_name = file_name[-1]

            msg.add_attachment(
                file_data,
                maintype="application",
                subtype="octet-stream",
                filename=file_name,
            )

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login(gmail_user, gmail_password)
        smtp.send_message(msg)
        print("Email sent!")


def do_something(filepath, download_path):
    try:
        new_df = pd.read_excel(filepath + '/NEW_FILE.xlsx', engine="openpyxl")
        #new_df = pd.read_excel('NEW_FILE.xlsx', engine="openpyxl")
        new_df.head(5)
        old_df = pd.read_excel(filepath + '/OLD_FILE.xlsx', engine="openpyxl")
        old_df.head(5)
        diff_df = pd.concat([new_df, old_df]).drop_duplicates(keep=False)
        if diff_df.empty:
            return

        diff_df['Report Date'] = diff_df['Report Date'].dt.strftime('%d/%m/%Y')
        diff_df['Occurrence No.'] = diff_df['Occurrence No.'].astype('int64')

        diff_df['Occurrence No.'] = diff_df['Occurrence No.'].astype('string')
        diff_df['Occurrence No.'] = diff_df['Occurrence No.'].str.zfill(3)
        new_df.reset_index(drop=True)

        car = ['Project No.', 'Project Description', 'Report Date',
            'Name', 'Occurrence No.', 'Description.', 'Date Closed.']
        fsr = ['Project No.', 'Project Description', 'Report Date',
            'Name', 'Occurrence No.', 'Description.']
        ncr = ['Project No.', 'Project Description',
            'Occurrence No.', 'Description.']
        par = ['Project No.', 'Project Description', 'Report Date',
            'Name', 'Occurrence No.', 'Description.', 'Date Closed.']
        qsr = ['Project No.', 'Project Description', 'Report Date',
            'Name', 'Occurrence No.', 'Discipline', 'Description.']

        i = diff_df.shape[0]
        for x in range(i):
            pre = '{}'

            file_name = diff_df['Type of occurrence:'].values[x]

            company = diff_df['Business Unit'].values[x]

            if file_name == 'CAR':
                output = [diff_df[label].values[x]
                        for label in car if label in diff_df]
            elif file_name == 'FSR':
                output = [diff_df[label].values[x]
                        for label in fsr if label in diff_df]
            elif file_name == 'NCR':
                output = [diff_df[label].values[x]
                        for label in ncr if label in diff_df]
            elif file_name == 'PAR':
                output = [diff_df[label].values[x]
                        for label in par if label in diff_df]
            else:
                output = [diff_df[label].values[x]
                        for label in qsr if label in diff_df]

            if company == 'MROTA':
                document = Document(file_name+'_Quinn.docx')
            else:
                document = Document(file_name+'.docx')
            folder_name = file_name
            length = len(output)
            print(length)
            tables = document.tables

            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if len(output) != 0:
                                to = str(output[0])
                                if pre in paragraph.text:
                                    paragraph.text = paragraph.text.replace(
                                        pre, to)
                                    output.pop(0)
                                    print(paragraph.text)

                            else:
                                to = ''
                                paragraph.text = paragraph.text.replace(pre, to)
            projnum = diff_df['Project No.'].values[x]
            occurrence_type = diff_df['Type of occurrence:'].values[x]
            occurrence_num = diff_df['Occurrence No.'].values[x]
            newfilename_w = str(projnum) + '-' \
                + str(occurrence_type) + '-' \
                + str(occurrence_num) + '.docx'

            download_file_path = os.path.join(  
                download_path, folder_name)

            if not os.path.exists(download_file_path):
                # create folder in the download folder
                os.makedirs(download_file_path)

            download_file_path_word = os.path.join(
                download_file_path, newfilename_w)

            document.save(download_file_path_word)

            convert(download_file_path_word)
            newfilename_pdf = str(projnum) + '-' \
                + str(occurrence_type) + '-' \
                + str(occurrence_num) + '.pdf'

            download_file_path_pdf = os.path.join(
                download_file_path, newfilename_pdf)

            download_file_path_word = [r"{}".format(    
                download_file_path_word)]
            download_file_path_pdf = [r"{}".format(
                download_file_path_pdf)] 
            download_file_path_pdf.extend(
                download_file_path_word)# extend to attach both pdf and word in email

            #email_to = [diff_df['Email'].values[x],'Rob.Williamson@graham.ca','Mike.Hudgins','Tim.waines@graham.ca']
            email_to = [diff_df['Email'].values[x]]
            subject = str(diff_df['Type of occurrence:'].values[x]) + ' File'
            body = f"Hey {diff_df['Name'].values[x]}\n\nPlease find attached word document and pdf of {file_name} ({diff_df['Project No.'].values[x]} / {diff_df['Occurrence No.'].values[x]})\n\nCheers\nQuality Rocks"

            email_auto_script(email_to, subject, body, download_file_path_pdf)
            raise FileNotFoundError

    except FileNotFoundError:
        print("An Invalid entry was found. Please Try again")
        email_to = ['divya.lakkad@graham.ca']
        rowindex = diff_df.index[x]
        subject = 'Invalid Entry in the quality document'
        body = f"Hey Divya \n\n There is an error in row {rowindex} in the NEW_FILE \n\n Thank you."
        exceptfilepath = ''
        email_auto_script(email_to, subject, body, exceptfilepath)
        
    now = datetime.datetime.now()
    time = now.strftime("%H_%M_%S")
    date = now.strftime("%d_%m_%Y")
    oldname = filepath + r'\OLD_FILE.xlsx'
    oldnamewithdate = filepath + r'\OLD_FILE_' + \
        str(date)+'_' + str(time) + '.xlsx'
    os.rename(oldname, oldnamewithdate)
    newname = filepath + r'\NEW_FILE.xlsx'
    shutil.copy(newname, oldname)


def do_loop(sc):
    filepath = r'C:\Users\divyal\Desktop\projects\Info_file'
    download_path = r'C:\Users\divyal\Desktop\projects\Quality_form_downlaod'
    do_something(filepath, download_path)
    print("NExt run in a 5 minutes")
    s.enter(300, 1, do_loop, (sc,))


s.enter(300, 1, do_loop, (s,))
s.run()
