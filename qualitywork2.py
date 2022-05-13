
import pandas as pd
import numpy
numpy.datetime_as_string
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches
import os
from datetime import date, timedelta
import datetime
from email.message import EmailMessage
import smtplib 
import base64
import email
import imaplib
from email.mime.multipart import MIMEMultipart 
from email.mime.text import MIMEText 
from email.mime.base import MIMEBase 
from email import encoders 
import email
import xlsxwriter
import sched, time
from docx import Document
from docx2pdf import convert

s = sched.scheduler(time.time, time.sleep)

def email_auto_script(to_email, subject, body, file_path):
    print('Sending Email!!!')

    gmail_user = "graham.scripting@gmail.com"
    gmail_password = "directorchris"

    
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = gmail_user
    msg["To"] = to_email
    #msg["Bcc"] = 'divya.lakkad@graham.ca'
    msg.set_content(body)

    for file in file_path:
        print(file)
        with open(file, "rb") as f:
            file_data = f.read()
            file_name = os.path.split(file)
            file_name = file_name[-1]

            msg.add_attachment(
                file_data,
                maintype="application",
                subtype="octet-stream",
                filename=file_name,
            )

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login(gmail_user, gmail_password)
        smtp.send_message(msg)
        print("Email sent!")

def do_something(sc): 

    new_list = []
    new_df = pd.read_excel('NEW_FILE.xlsx', engine="openpyxl")
    new_df.head(5)


    # %%
    old_df = pd.read_excel('OLD_FILE.xlsx',engine="openpyxl")


    # %%
    old_df.head(5)


    # %%
    diff_df = pd.concat([new_df,old_df]).drop_duplicates(keep=False)

    diff_df['Report Date'] = diff_df['Report Date'].dt.strftime('%d/%m/%Y')
    diff_df['Occurrence No.'] = diff_df['Occurrence No.'].astype('int64')

    diff_df['Occurrence No.'] = diff_df['Occurrence No.'].astype('string')
    diff_df['Occurrence No.'] = diff_df['Occurrence No.'].str.zfill(3)
    # %%
    new_df.reset_index(drop=True)


    # %%
    car = ['Project No.','Project Description','Report Date','Name','Occurrence No.','Description.']
    fsr = ['Project No.','Project Description','Report Date','Name','Occurrence No.','Description.']
    ncr = ['Project No.','Project Description','Occurrence No.','Description.']
    par = ['Project No.','Project Description','Report Date','Name','Occurrence No.','Description.']
    qsr = ['Project No.','Project Description','Report Date','Name','Occurrence No.','Type','Discipline','Description.']


    # %%
    CAR= [diff_df[label].values[0] for label in car if label in diff_df]
    FSR= [diff_df[label].values[0] for label in fsr if label in diff_df]
    NCR= [diff_df[label].values[0] for label in ncr if label in diff_df]
    PAR= [diff_df[label].values[0] for label in par if label in diff_df]
    QSR= [diff_df[label].values[0] for label in qsr if label in diff_df]
    print(CAR)


    # %%
    import docx

    def getText(filename):
        doc = docx.Document(filename)   
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return '\n'.join(fullText)


    # %%
    i = diff_df.shape[0]
    for x in range(i):
        print(x)

    files_email_path =[]
    i = diff_df.shape[0]
    for x in range(i):
        pre = '{}'
        CAR1= [diff_df[label].values[x] for label in car if label in diff_df]
        FSR1= [diff_df[label].values[x] for label in fsr if label in diff_df]
        NCR1= [diff_df[label].values[x] for label in ncr if label in diff_df]
        PAR1= [diff_df[label].values[x] for label in par if label in diff_df]
        QSR1= [diff_df[label].values[x] for label in qsr if label in diff_df]

        file_name= diff_df['Type of occurrence:'].values[x]
        company = diff_df['Business Unit'].values[x]
        
        if file_name == 'CAR':
            output = [diff_df[label].values[x] for label in car if label in diff_df]
        elif file_name == 'FSR':
            output = [diff_df[label].values[x] for label in fsr if label in diff_df]
        elif file_name == 'NCR':
            output = [diff_df[label].values[x] for label in ncr if label in diff_df]
        elif file_name == 'PAR':
            output = [diff_df[label].values[x] for label in par if label in diff_df]
        else:
            output = [diff_df[label].values[x] for label in qsr if label in diff_df]

        if company == 'MROTA':
            document = Document(file_name+'_Quinn.docx')
        else:
            document = Document(file_name+'.docx')
        
        length = len(output)
        print(length)
        tables = document.tables

        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs: 
                        if len(output) != 0:
                            to = str(output[0])
                            if pre in paragraph.text:
                                paragraph.text = paragraph.text.replace(pre, to)
                                output.pop(0)
                                print(paragraph.text)
                            
                        else:
                            to = ''
                            paragraph.text = paragraph.text.replace(pre, to)


        newfilename_w =  str(diff_df['Project No.'].values[x])+ '-' + str(diff_df['Occurrence No.'].values[x]) + '.docx'
        
        
        document.save(newfilename_w) 
        convert(newfilename_w)
        newfilename = str(diff_df['Project No.'].values[x])+ '-' + str(diff_df['Occurrence No.'].values[x]) + '.pdf'
        newfilepath_w =  os.path.join(os.path.dirname(__file__), newfilename_w)
        newfilepath =  os.path.join(os.path.dirname(__file__), newfilename)
        newfilepath_w= [r"{}".format(newfilepath_w)]
        newfilepath= [r"{}".format(newfilepath)]
        newfilepath.extend(newfilepath_w)
        #email_to = [diff_df['Email'].values[x],'Rob.Williamson@graham.ca','Mike.Hudgins','Tim.waines@graham.ca']
        email_to = [diff_df['Email'].values[x]]
        subject = str(diff_df['Type of occurrence:'].values[x]) + ' File'
        body = f"Hey {diff_df['Name'].values[x]}\n\nPlease find attached word document and pdf of {file_name} ({diff_df['Project No.'].values[x]} / {diff_df['Occurrence No.'].values[x]})\n\nCheers\nQuality Rocks"
        #email_auto_script(email_to, subject, body, newfilepath).

    x = datetime.datetime.now()
    x = x.strftime('%d/%m/%Y')
    writer = pd.ExcelWriter('OLD_FILE_'+str(x) +'.xlsx', engine='xlsxwriter')
    old_df.to_excel(writer, sheet_name='Form1')
    writer.save()
    new_df.to_excel('OLD_FILE.xlsx')
    os.remove('NEW_FILE.xlsx')

    s.enter(60, 5, do_something, (sc,))

s.enter(60, 5, do_something, (s,))
s.run()